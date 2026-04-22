"""
Resume Builder Service — Generates ATS-friendly .docx resumes from profile data.

Pulls from UserProfile.extra_data["resume_profile"] (student-editable fields)
and UserProfile base fields (ERP auto-fill) to produce a clean, single-column
Word document optimized for ATS scanners.

Templates:
  - classic: Clean, professional, single-column. Ideal for mass-recruiters.
"""
import io
import logging
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from app import models

logger = logging.getLogger("acadmix.resume_builder")


# ═══════════════════════════════════════════════════════════════════════════════
# Style Constants
# ═══════════════════════════════════════════════════════════════════════════════

_FONT_NAME = "Calibri"
_NAME_SIZE = Pt(18)
_CONTACT_SIZE = Pt(9)
_SECTION_HEADING_SIZE = Pt(11)
_BODY_SIZE = Pt(10)
_SMALL_SIZE = Pt(9)

_COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)   # Near-black
_COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)    # Body text
_COLOR_LINK = RGBColor(0x0A, 0x66, 0xC2)    # LinkedIn blue
_COLOR_SECTION = RGBColor(0x0D, 0x47, 0x71)  # Dark teal


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _set_doc_margins(doc: Document):
    """Set narrow margins for maximum content density."""
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def _add_run(paragraph, text: str, size=None, bold=False, color=None, font_name=None):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name or _FONT_NAME
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return run


def _add_section_heading(doc: Document, title: str):
    """Add a section heading with bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    _add_run(p, title.upper(), size=_SECTION_HEADING_SIZE, bold=True, color=_COLOR_SECTION)

    # Add bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '0D4771',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_body_text(doc: Document, text: str, bold=False, italic=False, space_after=Pt(2)):
    """Add a body-level paragraph."""
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = space_after
    run = _add_run(p, text, size=_BODY_SIZE, color=_COLOR_GRAY)
    if bold:
        run.bold = True
    if italic:
        run.font.italic = True
    return p


def _add_bullet(doc: Document, text: str):
    """Add a bullet-pointed paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    _add_run(p, text, size=_BODY_SIZE, color=_COLOR_GRAY)
    return p


def _add_placeholder(doc: Document, hint: str):
    """Add an instructional placeholder for missing data."""
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = _add_run(p, f"[{hint}]", size=_SMALL_SIZE, color=RGBColor(0xAA, 0xAA, 0xAA))
    run.font.italic = True
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# Template: Classic
# ═══════════════════════════════════════════════════════════════════════════════

def _build_classic(data: Dict[str, Any]) -> Document:
    """Build a classic, ATS-friendly single-column resume."""
    doc = Document()
    _set_doc_margins(doc)

    personal = data.get("personal", {})
    education = data.get("education_history", [])
    skills = data.get("skills", {})
    projects = data.get("projects", [])
    experience = data.get("experience", [])
    certs = data.get("certifications", [])
    achievements = data.get("achievements", [])
    summary = data.get("summary", "")

    # ── Name ───────────────────────────────────────────────────
    name = personal.get("name", "").strip()
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.space_after = Pt(2)
    if name:
        _add_run(name_p, name.upper(), size=_NAME_SIZE, bold=True, color=_COLOR_DARK)
    else:
        _add_run(name_p, "[YOUR FULL NAME]", size=_NAME_SIZE, bold=True,
                 color=RGBColor(0xAA, 0xAA, 0xAA))

    # ── Contact line ───────────────────────────────────────────
    contact_parts = []
    if personal.get("email"):
        contact_parts.append(personal["email"])
    if personal.get("phone"):
        contact_parts.append(personal["phone"])
    if personal.get("location"):
        contact_parts.append(personal["location"])

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.space_before = Pt(0)
        contact_p.space_after = Pt(2)
        _add_run(contact_p, " | ".join(contact_parts), size=_CONTACT_SIZE, color=_COLOR_GRAY)

    # ── Links line ─────────────────────────────────────────────
    link_parts = []
    if personal.get("linkedin"):
        link_parts.append(personal["linkedin"])
    if personal.get("github"):
        link_parts.append(personal["github"])
    if personal.get("portfolio"):
        link_parts.append(personal["portfolio"])

    if link_parts:
        links_p = doc.add_paragraph()
        links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_p.space_before = Pt(0)
        links_p.space_after = Pt(4)
        _add_run(links_p, " | ".join(link_parts), size=_CONTACT_SIZE, color=_COLOR_LINK)

    # ── Professional Summary ───────────────────────────────────
    if summary and summary.strip():
        _add_section_heading(doc, "Professional Summary")
        _add_body_text(doc, summary.strip())
    else:
        _add_section_heading(doc, "Professional Summary")
        _add_placeholder(doc, "Write 2-3 sentences about your technical interests, strengths, and career goals")

    # ── Education ──────────────────────────────────────────────
    _add_section_heading(doc, "Education")

    # Current education (auto-filled from ERP)
    current_edu = personal.get("current_education")
    if current_edu:
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        _add_run(p, current_edu.get("degree", "B.Tech"), size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
        _add_run(p, f" — {current_edu.get('institution', '')}", size=_BODY_SIZE, color=_COLOR_GRAY)
        details = []
        if current_edu.get("branch"):
            details.append(current_edu["branch"])
        if current_edu.get("batch"):
            details.append(f"Batch {current_edu['batch']}")
        if current_edu.get("cgpa"):
            details.append(f"CGPA: {current_edu['cgpa']}")
        if details:
            det_p = doc.add_paragraph()
            det_p.space_before = Pt(0)
            det_p.space_after = Pt(4)
            _add_run(det_p, " | ".join(details), size=_SMALL_SIZE, color=_COLOR_GRAY)

    # Prior education entries
    if education:
        for edu in education:
            p = doc.add_paragraph()
            p.space_before = Pt(0)
            p.space_after = Pt(1)
            level = edu.get("level", "")
            school = edu.get("school", "")
            _add_run(p, level, size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
            if school:
                _add_run(p, f" — {school}", size=_BODY_SIZE, color=_COLOR_GRAY)

            sub_parts = []
            if edu.get("board"):
                sub_parts.append(edu["board"])
            if edu.get("year"):
                sub_parts.append(str(edu["year"]))
            if edu.get("percentage"):
                sub_parts.append(f"{edu['percentage']}%")
            if sub_parts:
                sub_p = doc.add_paragraph()
                sub_p.space_before = Pt(0)
                sub_p.space_after = Pt(4)
                _add_run(sub_p, " | ".join(sub_parts), size=_SMALL_SIZE, color=_COLOR_GRAY)
    elif not current_edu:
        _add_placeholder(doc, "Add your education details — degree, institution, CGPA, year of graduation")

    # ── Skills ─────────────────────────────────────────────────
    has_skills = any(skills.get(k) for k in ["languages", "frameworks", "tools", "databases"])
    if has_skills:
        _add_section_heading(doc, "Technical Skills")
        for category, label in [("languages", "Languages"), ("frameworks", "Frameworks"),
                                 ("tools", "Tools & Platforms"), ("databases", "Databases")]:
            items = skills.get(category, [])
            if items:
                p = doc.add_paragraph()
                p.space_before = Pt(0)
                p.space_after = Pt(2)
                _add_run(p, f"{label}: ", size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
                _add_run(p, ", ".join(items), size=_BODY_SIZE, color=_COLOR_GRAY)
    else:
        _add_section_heading(doc, "Technical Skills")
        _add_placeholder(doc, "Languages: Python, Java, JavaScript  |  Frameworks: React, FastAPI  |  Tools: Git, Docker")

    # ── Projects ───────────────────────────────────────────────
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(1)
            _add_run(p, proj.get("title", "Untitled Project"), size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
            tech = proj.get("tech_stack", "")
            if tech:
                _add_run(p, f" | {tech}", size=_SMALL_SIZE, color=_COLOR_GRAY)
            link = proj.get("link", "")
            if link:
                _add_run(p, f" | {link}", size=_SMALL_SIZE, color=_COLOR_LINK)

            bullets = proj.get("bullets", [])
            if bullets:
                for b in bullets:
                    if b.strip():
                        _add_bullet(doc, b.strip())
            else:
                _add_placeholder(doc, "Describe what you built, technologies used, and quantifiable impact")
    else:
        _add_section_heading(doc, "Projects")
        _add_placeholder(doc, "Add 2-3 projects — include tech stack, what you built, and measurable outcomes")

    # ── Experience ─────────────────────────────────────────────
    if experience:
        _add_section_heading(doc, "Experience")
        for exp in experience:
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(1)
            _add_run(p, exp.get("role", "Role"), size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
            company = exp.get("company", "")
            if company:
                _add_run(p, f" — {company}", size=_BODY_SIZE, color=_COLOR_GRAY)

            sub_parts = []
            if exp.get("duration"):
                sub_parts.append(exp["duration"])
            if exp.get("location"):
                sub_parts.append(exp["location"])
            if sub_parts:
                sub_p = doc.add_paragraph()
                sub_p.space_before = Pt(0)
                sub_p.space_after = Pt(1)
                _add_run(sub_p, " | ".join(sub_parts), size=_SMALL_SIZE, color=_COLOR_GRAY)

            bullets = exp.get("bullets", [])
            if bullets:
                for b in bullets:
                    if b.strip():
                        _add_bullet(doc, b.strip())
            else:
                _add_placeholder(doc, "Describe your responsibilities and achievements — use action verbs and metrics")

    # ── Certifications ─────────────────────────────────────────
    if certs:
        _add_section_heading(doc, "Certifications")
        for cert in certs:
            p = doc.add_paragraph()
            p.space_before = Pt(0)
            p.space_after = Pt(2)
            _add_run(p, cert.get("name", "Certification"), size=_BODY_SIZE, bold=True, color=_COLOR_DARK)
            issuer = cert.get("issuer", "")
            year = cert.get("year", "")
            suffix = []
            if issuer:
                suffix.append(issuer)
            if year:
                suffix.append(str(year))
            if suffix:
                _add_run(p, f" — {', '.join(suffix)}", size=_BODY_SIZE, color=_COLOR_GRAY)

    # ── Achievements ───────────────────────────────────────────
    if achievements:
        _add_section_heading(doc, "Achievements")
        for ach in achievements:
            text = ach if isinstance(ach, str) else ach.get("title", "")
            if text.strip():
                _add_bullet(doc, text.strip())

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════════

async def generate_docx(
    user: dict,
    session: AsyncSession,
    template: str = "classic",
) -> io.BytesIO:
    """
    Generate a resume .docx from the student's resume profile data.

    Returns an in-memory BytesIO buffer containing the .docx file.
    """
    # Load profile
    result = await session.execute(
        select(models.UserProfile).where(
            models.UserProfile.user_id == user["id"],
        )
    )
    profile = result.scalars().first()

    if not profile:
        raise ValueError("No profile found")

    extra = profile.extra_data or {}
    resume = extra.get("resume_profile", {})

    # Build the merged data object
    data = {
        "personal": {
            "name": profile.full_name or user.get("full_name", ""),
            "email": resume.get("email") or user.get("email", ""),
            "phone": resume.get("phone") or getattr(profile, "phone", "") or "",
            "location": resume.get("location", ""),
            "linkedin": resume.get("linkedin", ""),
            "github": resume.get("github", ""),
            "portfolio": resume.get("portfolio", ""),
            "current_education": {
                "degree": "B.Tech",  # Default for now
                "institution": getattr(profile, "college_name", "") or profile.institution_id or "",
                "branch": profile.department or "",
                "batch": profile.batch or "",
            },
        },
        "summary": resume.get("summary", ""),
        "education_history": resume.get("education_history", []),
        "skills": resume.get("skills", {}),
        "projects": resume.get("projects", []),
        "experience": resume.get("experience", []),
        "certifications": resume.get("certifications", []),
        "achievements": resume.get("achievements", []),
    }

    # Generate document based on template
    if template == "classic":
        doc = _build_classic(data)
    else:
        doc = _build_classic(data)  # Fallback to classic

    # Write to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    student_name = (profile.full_name or "resume").replace(" ", "_")
    logger.info("Generated %s resume for student %s", template, user["id"])

    return buffer, f"{student_name}_Resume.docx"
